"""
DOCX Document Agent — full lifecycle for Word documents.

Uses python-docx for high-level paragraph/table access and lxml for
deep OOXML manipulation (tracked changes, comments, custom XML parts).
"""

from __future__ import annotations

import logging
import re
import uuid
from io import BytesIO
from typing import Any, Optional

from .base import (
    DocumentAgent,
    AgentCapability,
    CanvasBundle,
    CanvasNode,
    ComplianceConstraints,
    AtomGroup,
    EditOperation,
    ExportResult,
)

logger = logging.getLogger(__name__)

EMU_PER_POINT = 12700
EMU_PER_INCH = 914400
POINTS_PER_INCH = 72

_CATEGORY_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"past\s*perf", re.I), "past_performance"),
    (re.compile(r"key\s*personnel|team|staff", re.I), "key_personnel"),
    (re.compile(r"bio|resume|cv|curriculum", re.I), "key_personnel"),
    (re.compile(r"techni(?:cal)?\s*(?:approach|volume|narrative)", re.I), "technical_approach"),
    (re.compile(r"management\s*(?:approach|plan|volume)", re.I), "management_approach"),
    (re.compile(r"cost|budget|pricing", re.I), "cost_volume"),
    (re.compile(r"capabilit", re.I), "capability_statement"),
    (re.compile(r"commerciali[sz]ation", re.I), "commercialization"),
    (re.compile(r"abstract|summary|overview|introduction", re.I), "abstract"),
    (re.compile(r"qualif|experience|corporate", re.I), "qualifications"),
    (re.compile(r"schedule|timeline|milestone", re.I), "schedule"),
    (re.compile(r"risk|mitigation", re.I), "risk_management"),
    (re.compile(r"quality|assurance", re.I), "quality"),
    (re.compile(r"facil|lab|equipment", re.I), "facilities"),
    (re.compile(r"subcontract|teaming|partner", re.I), "teaming"),
    (re.compile(r"security|clearance|itar", re.I), "security"),
    (re.compile(r"transition|sustainment", re.I), "transition_plan"),
    (re.compile(r"data\s*(?:management|rights)|intellectual", re.I), "data_rights"),
]


def _infer_category(text: str) -> tuple[str, float]:
    for pattern, cat in _CATEGORY_PATTERNS:
        if pattern.search(text):
            return cat, 0.8
    return "general", 0.3


def _emu_to_pt(emu: int | None) -> float:
    if emu is None:
        return 72.0
    return emu / EMU_PER_POINT


def _make_node(
    node_type: str,
    content: dict[str, Any],
    source: str = "imported",
) -> CanvasNode:
    return CanvasNode(
        id=str(uuid.uuid4()),
        type=node_type,
        content=content,
        provenance={"source": source, "drafted_by": "system:docx_agent"},
        history=[{
            "actor_id": "system:docx_agent",
            "actor_name": "DOCX Agent",
            "action": "created",
        }],
    )


class DocxAgent(DocumentAgent):
    @property
    def format_id(self) -> str:
        return "docx"

    @property
    def display_name(self) -> str:
        return "Word Document Agent"

    @property
    def file_extensions(self) -> list[str]:
        return ["docx", "doc"]

    @property
    def capabilities(self) -> set[AgentCapability]:
        return {
            AgentCapability.READ_NATIVE,
            AgentCapability.WRITE_NATIVE,
            AgentCapability.TRACKED_CHANGES,
            AgentCapability.COMMENTS,
            AgentCapability.HEADERS_FOOTERS,
            AgentCapability.PAGE_NUMBERS,
            AgentCapability.TABLE_OF_CONTENTS,
            AgentCapability.IMAGES,
            AgentCapability.WATERMARKS,
            AgentCapability.MERGE_FIELDS,
            AgentCapability.PDF_RENDER,
        }

    # ── INGEST ─────────────────────────────────────────────────────────

    async def ingest(self, file_bytes: bytes, filename: str) -> CanvasBundle:
        from docx import Document

        doc = Document(BytesIO(file_bytes))
        nodes: list[CanvasNode] = []
        list_buffer: list[dict[str, Any]] = []
        list_type: str | None = None

        def flush_list():
            nonlocal list_buffer, list_type
            if not list_buffer:
                return
            nodes.append(_make_node(
                list_type or "bulleted_list",
                {"items": list_buffer},
            ))
            list_buffer = []
            list_type = None

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()

            # Heading detection
            heading_match = re.match(r"heading\s*(\d)", style_name)
            if heading_match:
                flush_list()
                level = min(int(heading_match.group(1)), 3)
                nodes.append(_make_node("heading", {
                    "level": level,
                    "text": para.text.strip(),
                }))
                continue

            # List detection
            if "list" in style_name:
                is_numbered = "number" in style_name or "ordered" in style_name
                lt = "numbered_list" if is_numbered else "bulleted_list"

                if list_type and list_type != lt:
                    flush_list()

                list_type = lt
                list_buffer.append({
                    "text": para.text.strip(),
                    "indent_level": para.paragraph_format.left_indent or 0,
                })
                continue

            flush_list()

            text = para.text.strip()
            if not text:
                continue

            # Build inline formats from runs
            inline_formats = []
            offset = 0
            for run in para.runs:
                run_len = len(run.text)
                if run_len == 0:
                    continue
                if run.bold:
                    inline_formats.append({"start": offset, "length": run_len, "format": "bold"})
                if run.italic:
                    inline_formats.append({"start": offset, "length": run_len, "format": "italic"})
                if run.underline:
                    inline_formats.append({"start": offset, "length": run_len, "format": "underline"})
                if run.font.superscript:
                    inline_formats.append({"start": offset, "length": run_len, "format": "superscript"})
                if run.font.subscript:
                    inline_formats.append({"start": offset, "length": run_len, "format": "subscript"})
                offset += run_len

            content: dict[str, Any] = {"text": text}
            if inline_formats:
                content["inline_formats"] = inline_formats

            nodes.append(_make_node("text_block", content))

        flush_list()

        # Tables
        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])

            if rows_data:
                headers = rows_data[0] if rows_data else []
                body = rows_data[1:] if len(rows_data) > 1 else []
                nodes.append(_make_node("table", {
                    "headers": headers,
                    "rows": body,
                }))

        # Metadata
        cp = doc.core_properties
        metadata: dict[str, Any] = {
            "title": cp.title or "",
            "author": cp.author or "",
            "subject": cp.subject or "",
            "keywords": cp.keywords or "",
            "source_filename": filename,
        }
        if cp.created:
            metadata["created"] = cp.created.isoformat()
        if cp.modified:
            metadata["modified"] = cp.modified.isoformat()

        # Constraints from first section
        constraints = ComplianceConstraints()
        if doc.sections:
            sec = doc.sections[0]
            constraints.margins = {
                "top": _emu_to_pt(sec.top_margin),
                "right": _emu_to_pt(sec.right_margin),
                "bottom": _emu_to_pt(sec.bottom_margin),
                "left": _emu_to_pt(sec.left_margin),
            }
            if sec.page_width and sec.page_height:
                w_pt = _emu_to_pt(sec.page_width)
                h_pt = _emu_to_pt(sec.page_height)
                metadata["page_width_pt"] = w_pt
                metadata["page_height_pt"] = h_pt

            # Font defaults from first paragraph's style
            if doc.paragraphs:
                first_font = doc.paragraphs[0].style.font
                if first_font and first_font.name:
                    constraints.font_family = first_font.name
                if first_font and first_font.size:
                    constraints.font_size = first_font.size / EMU_PER_POINT

        return CanvasBundle(
            document_id=str(uuid.uuid4()),
            nodes=nodes,
            constraints=constraints,
            metadata=metadata,
            source_format="docx",
            source_agent=self.display_name,
        )

    # ── ATOMIZE ────────────────────────────────────────────────────────

    async def atomize(self, bundle: CanvasBundle) -> list[AtomGroup]:
        if not bundle.nodes:
            return []

        atoms: list[AtomGroup] = []
        group: list[CanvasNode] = []
        heading_text: str | None = None
        heading_level = 0
        char_offset = 0

        def flush():
            nonlocal group, heading_text, char_offset
            if not group:
                return
            text = " ".join(_node_text(n) for n in group)
            cat, conf = _infer_category(heading_text or text[:500])
            tags = [cat]
            if heading_text:
                tags.append(f"heading:{heading_text[:80]}")
            fn = bundle.metadata.get("source_filename", "")
            if fn:
                tags.append(f"source:{fn[:50]}")

            atoms.append(AtomGroup(
                nodes=list(group),
                heading_text=heading_text,
                suggested_category=cat,
                suggested_tags=tags,
                confidence=conf,
                char_offset=char_offset,
                char_length=len(text),
            ))
            char_offset += len(text)
            group = []
            heading_text = None

        for node in bundle.nodes:
            if node.type == "heading":
                level = node.content.get("level", 1)
                if group and level <= heading_level:
                    flush()
                if not group:
                    heading_text = node.content.get("text")
                    heading_level = level
                group.append(node)
            else:
                if not group:
                    heading_text = None
                    heading_level = 0
                group.append(node)

        flush()

        # Single huge atom → split by paragraph count
        if len(atoms) == 1 and len(atoms[0].nodes) > 8:
            return self._split_by_count(atoms[0].nodes, bundle, 4)

        return atoms

    def _split_by_count(
        self, nodes: list[CanvasNode], bundle: CanvasBundle, size: int
    ) -> list[AtomGroup]:
        atoms: list[AtomGroup] = []
        group: list[CanvasNode] = []
        count = 0
        offset = 0
        fn = bundle.metadata.get("source_filename", "")

        for node in nodes:
            group.append(node)
            if node.type == "text_block":
                count += 1
            if count >= size:
                text = " ".join(_node_text(n) for n in group)
                cat, conf = _infer_category(text[:500])
                atoms.append(AtomGroup(
                    nodes=list(group),
                    heading_text=None,
                    suggested_category=cat,
                    suggested_tags=[cat, f"source:{fn[:50]}"],
                    confidence=conf,
                    char_offset=offset,
                    char_length=len(text),
                ))
                offset += len(text)
                group = []
                count = 0

        if group:
            text = " ".join(_node_text(n) for n in group)
            cat, conf = _infer_category(text[:500])
            atoms.append(AtomGroup(
                nodes=list(group),
                heading_text=None,
                suggested_category=cat,
                suggested_tags=[cat, f"source:{fn[:50]}"],
                confidence=conf,
                char_offset=offset,
                char_length=len(text),
            ))

        return atoms

    # ── EXPORT ─────────────────────────────────────────────────────────

    async def export(self, bundle: CanvasBundle) -> ExportResult:
        from docx import Document
        from docx.shared import Pt, Inches, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        c = bundle.constraints

        # Page setup
        if doc.sections:
            sec = doc.sections[0]
            if c.margins:
                sec.top_margin = int(c.margins.get("top", 72) * EMU_PER_POINT)
                sec.right_margin = int(c.margins.get("right", 72) * EMU_PER_POINT)
                sec.bottom_margin = int(c.margins.get("bottom", 72) * EMU_PER_POINT)
                sec.left_margin = int(c.margins.get("left", 72) * EMU_PER_POINT)

        # Default font
        style = doc.styles["Normal"]
        if c.font_family:
            style.font.name = c.font_family
        if c.font_size:
            style.font.size = Pt(c.font_size)

        # Header
        if c.header_template and doc.sections:
            header = doc.sections[0].header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            text = _interpolate(c.header_template, bundle.variables)
            p.text = text
            if c.font_size:
                for run in p.runs:
                    run.font.size = Pt(max(c.font_size - 2, 8))

        # Footer with page numbers
        if c.footer_template and doc.sections:
            footer = doc.sections[0].footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            text = _interpolate(c.footer_template, bundle.variables)
            # Replace {n} and {N} with Word field codes
            if "{n}" in text or "{N}" in text:
                parts = text.replace("{N}", "{NUMPAGES}").replace("{n}", "{PAGE}").split("{")
                p.clear()
                for i, part in enumerate(parts):
                    if i == 0:
                        p.add_run(part)
                    elif part.startswith("PAGE}"):
                        _add_page_number_field(p)
                        p.add_run(part[5:])
                    elif part.startswith("NUMPAGES}"):
                        _add_num_pages_field(p)
                        p.add_run(part[9:])
                    else:
                        p.add_run("{" + part)
            else:
                p.text = text

        # Render nodes
        for node in bundle.nodes:
            if node.type == "heading":
                level = node.content.get("level", 1)
                doc.add_heading(node.content.get("text", ""), level=level)

            elif node.type == "text_block":
                p = doc.add_paragraph()
                text = node.content.get("text", "")
                formats = node.content.get("inline_formats", [])

                if formats:
                    _render_formatted_paragraph(p, text, formats, c)
                else:
                    run = p.add_run(text)
                    if c.font_family:
                        run.font.name = c.font_family
                    if c.font_size:
                        run.font.size = Pt(c.font_size)

            elif node.type in ("bulleted_list", "numbered_list"):
                style_name = "List Bullet" if node.type == "bulleted_list" else "List Number"
                for item in node.content.get("items", []):
                    p = doc.add_paragraph(
                        item.get("text", ""),
                        style=style_name,
                    )

            elif node.type == "table":
                headers = node.content.get("headers", [])
                rows = node.content.get("rows", [])
                col_count = max(len(headers), max((len(r) for r in rows), default=0))
                if col_count == 0:
                    continue

                tbl = doc.add_table(
                    rows=1 + len(rows),
                    cols=col_count,
                    style="Table Grid",
                )

                # Headers
                for j, h in enumerate(headers):
                    cell = tbl.rows[0].cells[j]
                    cell.text = h if isinstance(h, str) else h.get("text", "")
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

                # Data rows
                for i, row in enumerate(rows):
                    for j, val in enumerate(row):
                        cell = tbl.rows[i + 1].cells[j]
                        cell.text = val if isinstance(val, str) else val.get("text", "")

            elif node.type == "page_break":
                doc.add_page_break()

        # Watermark
        if c.watermark:
            _add_watermark(doc, c.watermark)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        fn = bundle.metadata.get("source_filename", "document")
        base = fn.rsplit(".", 1)[0] if "." in fn else fn

        return ExportResult(
            file_bytes=buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{base}.docx",
            format="docx",
        )

    async def export_pdf(self, bundle: CanvasBundle) -> ExportResult:
        native = await self.export(bundle)
        return await self._convert_to_pdf(native)


# ── Helpers ────────────────────────────────────────────────────────────


def _node_text(node: CanvasNode) -> str:
    c = node.content
    if node.type in ("heading", "text_block"):
        return c.get("text", "")
    if node.type in ("bulleted_list", "numbered_list"):
        return " ".join(item.get("text", "") for item in c.get("items", []))
    if node.type == "table":
        parts = [
            h if isinstance(h, str) else h.get("text", "")
            for h in c.get("headers", [])
        ]
        for row in c.get("rows", []):
            parts.extend(
                v if isinstance(v, str) else v.get("text", "")
                for v in row
            )
        return " ".join(parts)
    return ""


def _interpolate(template: str, variables: dict[str, str]) -> str:
    result = template
    for key, val in variables.items():
        result = result.replace(f"{{{key}}}", val)
    return result


def _render_formatted_paragraph(p, text: str, formats: list[dict], constraints):
    from docx.shared import Pt

    events: list[tuple[int, str, str]] = []
    for f in formats:
        s, l, fmt = f["start"], f["length"], f["format"]
        events.append((s, "open", fmt))
        events.append((s + l, "close", fmt))
    events.sort(key=lambda e: (e[0], 0 if e[1] == "close" else 1))

    breaks: list[int] = sorted({0, len(text)} | {e[0] for e in events})
    active: set[str] = set()
    pos = 0

    for b in breaks:
        if b > pos and b <= len(text):
            chunk = text[pos:b]
            run = p.add_run(chunk)
            if constraints.font_family:
                run.font.name = constraints.font_family
            if constraints.font_size:
                run.font.size = Pt(constraints.font_size)
            run.bold = "bold" in active
            run.italic = "italic" in active
            run.underline = "underline" in active
            run.font.superscript = "superscript" in active
            run.font.subscript = "subscript" in active

        for ev in events:
            if ev[0] == b:
                if ev[1] == "open":
                    active.add(ev[2])
                else:
                    active.discard(ev[2])
        pos = b


def _add_page_number_field(paragraph):
    from docx.oxml.ns import qn
    from lxml import etree

    run = paragraph.add_run()
    fldChar = etree.SubElement(run._r, qn("w:fldChar"))
    fldChar.set(qn("w:fldCharType"), "begin")

    run2 = paragraph.add_run()
    instrText = etree.SubElement(run2._r, qn("w:instrText"))
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "

    run3 = paragraph.add_run()
    fldChar2 = etree.SubElement(run3._r, qn("w:fldChar"))
    fldChar2.set(qn("w:fldCharType"), "end")


def _add_num_pages_field(paragraph):
    from docx.oxml.ns import qn
    from lxml import etree

    run = paragraph.add_run()
    fldChar = etree.SubElement(run._r, qn("w:fldChar"))
    fldChar.set(qn("w:fldCharType"), "begin")

    run2 = paragraph.add_run()
    instrText = etree.SubElement(run2._r, qn("w:instrText"))
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " NUMPAGES "

    run3 = paragraph.add_run()
    fldChar2 = etree.SubElement(run3._r, qn("w:fldChar"))
    fldChar2.set(qn("w:fldCharType"), "end")


def _add_watermark(doc, text: str):
    """Add a diagonal text watermark to the document header."""
    from docx.oxml.ns import qn
    from lxml import etree

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        r = p.add_run()

        # VML shape for watermark
        vml = (
            f'<v:shapetype id="_x0000_t136" coordsize="21600,21600" '
            f'o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e" '
            f'xmlns:v="urn:schemas-microsoft-com:vml" '
            f'xmlns:o="urn:schemas-microsoft-com:office:office">'
            f'</v:shapetype>'
            f'<v:shape id="PowerPlusWaterMarkObject" '
            f'o:spid="_x0000_s2049" type="#_x0000_t136" '
            f'style="position:absolute;margin-left:0;margin-top:0;'
            f'width:500pt;height:100pt;rotation:315;z-index:-251657216;'
            f'mso-position-horizontal:center;mso-position-vertical:center" '
            f'o:allowincell="f" fillcolor="silver" stroked="f" '
            f'xmlns:v="urn:schemas-microsoft-com:vml" '
            f'xmlns:o="urn:schemas-microsoft-com:office:office">'
            f'<v:fill opacity=".25"/>'
            f'<v:textpath style="font-family:&quot;Arial&quot;;font-size:72pt" '
            f'string="{text}"/>'
            f'</v:shape>'
        )

        try:
            shape_elem = etree.fromstring(
                f'<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                f'{vml}</w:pict>'
            )
            r._r.append(shape_elem)
        except Exception:
            logger.warning("Failed to add watermark to DOCX")
